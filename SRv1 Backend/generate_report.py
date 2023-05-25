import pandas as pd
from docx import Document
from sqlalchemy import create_engine

def create_regulatory_report():
    engine = create_engine('mysql://user:pass@localhost/dbname')

    sql_query = """
    SELECT user.name AS user, transaction.value AS value, transaction.date AS date
    FROM user
    JOIN transaction ON user.id = transaction.user_id
    WHERE transaction.date BETWEEN '2022-01-01' AND '2022-12-31'
    ORDER BY transaction.date
    """

    df = pd.read_sql_query(sql_query, engine)

    document = Document()

    document.add_heading('Regulatory Report', 0)
    document.add_paragraph('Transactions from 2022:')

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'User'
    hdr_cells[1].text = 'Value'
    hdr_cells[2].text = 'Date'

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['user'])
        row_cells[1].text = str(row['value'])
        row_cells[2].text = str(row['date'])

    document.save('regulatory_report.docx')

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def create_pdf_report(data, filename):
    doc = SimpleDocTemplate(filename)
    elements = []
    styles = getSampleStyleSheet()

    title = Paragraph("My Report", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 0.2*inch))

    for item in data:
        ptext = '<font size=12>%s</font>' % item.strip()
        elements.append(Paragraph(ptext, styles["Normal"]))
        elements.append(Spacer(1, 0.2 * inch))

    doc.build(elements)

def create_docx_report(data, filename):
    doc = Document()

    doc.add_heading('My Report', 0)

    for item in data:
        doc.add_paragraph(item.strip())

    doc.save(filename)
